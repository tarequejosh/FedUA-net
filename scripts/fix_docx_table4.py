"""
Fix Table 4 (scarcity table) in paper_final.docx.
Pandoc generated it as 1-row/2-col; we rebuild it entirely.
"""
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document("paper_final.docx")

# Verify tables
for i, t in enumerate(doc.tables):
    print(f"Table {i+1}: {len(t.rows)}r x {len(t.rows[0].cells)}c, first cell: {t.rows[0].cells[0].text[:30]!r}")

T4_DATA = [
    ["Method",               "N=546 (Full)",     "N=200 (Scarcity)", "N=100 (Extreme)"],
    ["Local-Only",           "84.62 \u00b1 0.74", "79.77 \u00b1 1.90", "81.20 \u00b1 5.20"],
    ["FedBN",                "85.47 \u00b1 2.07", "73.22 \u00b1 1.31", "66.67 \u00b1 1.48"],
    ["FedUA-Net (Uniform)",  "88.60 \u00b1 3.56", "77.21 \u00b1 4.71", "70.09 \u00b1 4.52"],
]

def rebuild_table(doc, old_table_idx, new_data):
    """Replace the old table (by index) with a fresh one."""
    old_tbl = doc.tables[old_table_idx]
    old_tbl_element = old_tbl._element
    parent = old_tbl_element.getparent()
    idx = list(parent).index(old_tbl_element)

    # Create new table
    new_tbl = doc.add_table(rows=len(new_data), cols=len(new_data[0]))
    new_tbl.style = 'Table Grid'

    for ri, row_data in enumerate(new_data):
        for ci, text in enumerate(row_data):
            cell = new_tbl.rows[ri].cells[ci]
            para = cell.paragraphs[0]
            run = para.add_run(text)
            run.font.size = Pt(9)
            if ri == 0:
                run.bold = True
            if "FedUA-Net" in row_data[0] and ri > 0:
                run.bold = True
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER if ci > 0 else WD_ALIGN_PARAGRAPH.LEFT

    # Move the new table element to the old position
    new_tbl_element = new_tbl._element
    new_tbl_element.getparent().remove(new_tbl_element)
    parent.insert(idx, new_tbl_element)
    # Remove old table
    parent.remove(old_tbl_element)
    print(f"  Rebuilt Table {old_table_idx+1} with {len(new_data)} rows x {len(new_data[0])} cols")

rebuild_table(doc, 3, T4_DATA)  # Table 4 (0-indexed = 3)

doc.save("paper_final.docx")
print("Done. paper_final.docx updated.")
