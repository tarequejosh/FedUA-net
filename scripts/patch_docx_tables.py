"""
Patches paper_final.docx tables with correct numerical data from verified results.
Fills in all cells that pandoc left blank due to LaTeX math.
"""
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
import copy, re

doc = Document("paper_final.docx")

# ─────────────────────────────────────────────────────────
# TABLE 1 — Main results (Table I in paper)
# ─────────────────────────────────────────────────────────
# Columns: Method | Accuracy (%) | Macro F1 (%) | MCC | Raw ECE | Calibrated ECE | APS Set Size (α=0.10)
T1_DATA = [
    ["Method",                        "Accuracy (%)",   "Macro F1 (%)",   "MCC",           "Raw ECE", "Calibrated ECE", "APS Set Size (α=0.10)"],
    ["Local-Only",                    "94.01 ± 0.32",   "93.80 ± 0.48",   "0.906 ± 0.005", "0.0463",  "0.0251",         "2.01 ± 0.07"],
    ["Ditto",                         "93.93 ± 0.51",   "93.67 ± 0.62",   "0.904 ± 0.009", "0.0597",  "0.0289",         "2.18 ± 0.09"],
    ["FedUA-Net (CKA-Personalized)",  "93.87 ± 0.94",   "93.73 ± 1.08",   "0.904 ± 0.016", "0.0588",  "0.0307",         "2.19 ± 0.17"],
    ["Centralized (Pooled)",          "93.45 ± 1.30",   "93.36 ± 1.38",   "0.895 ± 0.024", "0.0714",  "—",              "—"],
    ["FedUA-Net (Uniform Baseline)",  "93.30 ± 1.13",   "93.05 ± 1.36",   "0.894 ± 0.019", "0.0504",  "0.0307",         "2.33 ± 0.23"],
    ["FedAvg",                        "92.36 ± 0.92",   "91.88 ± 1.23",   "0.878 ± 0.016", "0.0697",  "0.0398",         "2.22 ± 0.11"],
    ["FedBN",                         "92.34 ± 0.99",   "91.76 ± 1.02",   "0.877 ± 0.017", "0.0650",  "0.0379",         "2.20 ± 0.16"],
    ["FedBABU",                       "91.99 ± 0.35",   "91.50 ± 0.34",   "0.873 ± 0.006", "0.0640",  "0.1913",         "2.14 ± 0.17"],
    ["FedProx",                       "91.97 ± 1.03",   "91.46 ± 1.21",   "0.870 ± 0.019", "0.0680",  "0.0888",         "2.26 ± 0.18"],
]

# ─────────────────────────────────────────────────────────
# TABLE 2 — Per-client accuracy (Table II in paper)
# ─────────────────────────────────────────────────────────
# Columns: Method | Site A (MRI) | Site B (US) | Site C (X-Ray)
T2_DATA = [
    ["Method",                       "Site A (MRI)",    "Site B (US)",    "Site C (X-Ray)"],
    ["Local-Only",                   "96.21 ± 0.10",    "90.31 ± 0.99",   "95.50 ± 0.43"],
    ["Ditto",                        "96.15 ± 0.07",    "90.60 ± 2.26",   "95.03 ± 0.99"],
    ["FedUA-Net (CKA-Pers.)",        "95.89 ± 0.18",    "90.26 ± 3.00",   "95.46 ± 0.25"],
    ["Centralized",                  "96.38 ± 0.25",    "88.03 ± 4.52",   "95.93 ± 0.66"],
    ["FedUA-Net (Baseline)",         "96.06 ± 0.29",    "88.60 ± 3.56",   "95.23 ± 0.42"],
    ["FedAvg",                       "96.21 ± 0.18",    "85.47 ± 3.08",   "95.39 ± 0.72"],
    ["FedBN",                        "96.10 ± 0.25",    "85.47 ± 3.42",   "95.44 ± 0.57"],
    ["FedBABU",                      "96.10 ± 0.31",    "84.62 ± 1.48",   "95.24 ± 0.79"],
    ["FedProx",                      "96.13 ± 0.23",    "84.33 ± 3.56",   "95.45 ± 0.67"],
]

# ─────────────────────────────────────────────────────────
# TABLE 3 — Ablation (Table III in paper, 5 rows + header)
# ─────────────────────────────────────────────────────────
# Columns: Configuration | Attention | Local FT | Hosp. A | Hosp. B | Hosp. C | Mean Acc.
T3_DATA = [
    ["Configuration",               "Attention Module", "Local FT", "Hosp. A (MRI)",  "Hosp. B (US)",  "Hosp. C (X-Ray)", "Multi-Task Mean Acc. (%)"],
    ["FedBN Baseline",              "None",             "✗",        "96.02 ± 0.34",   "54.99 ± 3.00",  "95.65 ± 0.39",    "82.22 ± 0.96"],
    ["Base Personalization",        "None",             "✓",        "96.40 ± 0.31",   "82.91 ± 3.73",  "95.30 ± 0.60",    "91.53 ± 1.19"],
    ["+ Spatial Attention",         "Spatial",          "✓",        "95.98 ± 0.32",   "80.06 ± 2.15",  "95.56 ± 0.25",    "90.53 ± 0.68"],
    ["+ Channel Attention",         "Channel",          "✓",        "96.00 ± 0.25",   "81.48 ± 3.00",  "95.22 ± 0.22",    "90.90 ± 0.94"],
    ["FedUA-Net (Dual CBAM)",       "Dual CBAM",        "✓",        "96.29 ± 0.28",   "83.48 ± 4.04",  "95.50 ± 0.41",    "91.75 ± 1.34"],
]

# ─────────────────────────────────────────────────────────
# TABLE 4 — Scarcity (Table IV)
# ─────────────────────────────────────────────────────────
T4_DATA = [
    ["Method",                  "N=546 (100%)",   "N=200 (Scarcity)", "N=100 (Extreme)"],
    ["Local-Only",              "90.31 ± 0.99",   "84.33 ± 3.00",     "81.20 ± 5.20"],
    ["FedBN",                   "85.47 ± 3.42",   "73.22 ± 1.31",     "66.67 ± 1.48"],
    ["FedUA-Net (Uniform)",     "88.60 ± 3.56",   "77.21 ± 4.71",     "70.09 ± 4.52"],
]

# ─────────────────────────────────────────────────────────
# TABLE 5 — might be a stats table; skip if unrecognized
# ─────────────────────────────────────────────────────────

ALL_TABLES = [T1_DATA, T2_DATA, T3_DATA, T4_DATA]

def set_cell_text(cell, text, bold=False):
    """Replace cell content with plain text, optionally bold."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    if cell.paragraphs:
        para = cell.paragraphs[0]
    else:
        para = cell.add_paragraph()
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(9)

def patch_table(tbl, data):
    rows = tbl.rows
    n_rows = min(len(rows), len(data))
    for ri in range(n_rows):
        n_cols = min(len(rows[ri].cells), len(data[ri]))
        for ci in range(n_cols):
            is_header = (ri == 0)
            is_proposed = any(kw in data[ri][0] for kw in ["CKA-Personalized", "FedUA-Net (Dual"])
            txt = data[ri][ci]
            set_cell_text(rows[ri].cells[ci], txt, bold=(is_header or is_proposed))

print(f"Document has {len(doc.tables)} tables")
for i, (tbl, data) in enumerate(zip(doc.tables, ALL_TABLES)):
    print(f"Patching Table {i+1}: {len(tbl.rows)} rows x {len(tbl.rows[0].cells)} cols -> {len(data)} data rows")
    patch_table(tbl, data)

doc.save("paper_final.docx")
print("\nSaved paper_final.docx with patched tables.")
